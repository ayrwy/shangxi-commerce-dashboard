from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SHOTS = ROOT / "screenshots"
OUTPUT = ROOT / "商析_全栈电商数据分析平台_作品集.docx"

INK = "292721"
MUTED = "6F6A61"
PAPER = "FCFAF5"
CANVAS = "EEE9DF"
OLIVE = "737B59"
SIGNAL = "C65D3D"
LINE = "D8D0C3"


def set_font(run, size=None, color=None, bold=None, italic=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge not in kwargs:
            continue
        attrs = kwargs[edge]
        tag = qn(f"w:{edge}")
        element = tcBorders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def set_table_widths(table, widths):
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "9360")
    tblW.set(qn("w:type"), "dxa")
    tblLayout = tblPr.first_child_found_in("w:tblLayout")
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")
    grid = tbl.tblGrid
    for i, width in enumerate(widths):
        grid.gridCol_lst[i].set(qn("w:w"), str(width))
    for row in table.rows:
        for i, width in enumerate(widths):
            tcW = row.cells[i]._tc.tcPr.first_child_found_in("w:tcW")
            tcW.set(qn("w:w"), str(width))
            tcW.set(qn("w:type"), "dxa")


def add_rule(paragraph, color=LINE, sz="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_text(doc, text="", size=11, color=INK, bold=False, italic=False, align=None, before=0, after=6, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=16 if level == 1 else 12.5, color=INK if level == 1 else OLIVE, bold=True)
    return p


def add_bullet(doc, label, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(label)
    set_font(r, size=10.5, color=INK, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5, color=INK)
    return p


def add_figure(doc, filename, caption, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(str(SHOTS / filename), width=Inches(width))
    cap = add_text(doc, caption, size=9, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    return cap


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [2340, 2340, 2340, 2340])
    data = [
        ("前端", "React + TypeScript"),
        ("后端", "Django REST"),
        ("数据库", "MySQL"),
        ("数据规模", "最大明细表 < 5,000 行"),
    ]
    for cell, (label, value) in zip(table.rows[0].cells, data):
        set_cell_shading(cell, "F6F2EA")
        set_cell_margins(cell, 130, 150, 130, 150)
        set_cell_border(cell, top={"val": "single", "sz": "6", "color": LINE}, bottom={"val": "single", "sz": "6", "color": LINE}, left={"val": "single", "sz": "6", "color": LINE}, right={"val": "single", "sz": "6", "color": LINE})
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label)
        set_font(r, size=8.5, color=MUTED, bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        set_font(r2, size=10.5, color=INK, bold=True)
    return table


def add_architecture(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [3120, 3120, 3120])
    items = [
        ("01  数据入口", "CSV 多文件导入\n浏览器端解析与字段预览\n表角色与行粒度确认"),
        ("02  服务层", "Django 用户会话与权限\nREST API / CSRF / CORS\n导入任务与关系确认"),
        ("03  数据与分析", "MySQL 持久化业务数据\n指标引擎 / 漏斗 / RFM\n图表看板与数据诊断"),
    ]
    for cell, (title, body) in zip(table.rows[0].cells, items):
        set_cell_shading(cell, PAPER)
        set_cell_margins(cell, 160, 180, 160, 180)
        set_cell_border(cell, top={"val": "single", "sz": "10", "color": OLIVE}, bottom={"val": "single", "sz": "6", "color": LINE}, left={"val": "single", "sz": "6", "color": LINE}, right={"val": "single", "sz": "6", "color": LINE})
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(title)
        set_font(r, size=10, color=OLIVE, bold=True)
        for idx, line in enumerate(body.split("\n")):
            pp = cell.add_paragraph()
            pp.paragraph_format.space_after = Pt(3 if idx < 2 else 0)
            rr = pp.add_run(line)
            set_font(rr, size=9.5, color=INK)
    return table


def add_resume_block(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3E0D7")
    set_cell_margins(cell, 180, 220, 180, 220)
    set_cell_border(cell, top={"val": "single", "sz": "8", "color": SIGNAL}, bottom={"val": "single", "sz": "8", "color": SIGNAL}, left={"val": "single", "sz": "8", "color": SIGNAL}, right={"val": "single", "sz": "8", "color": SIGNAL})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("简历可用描述")
    set_font(r, size=11.5, color=SIGNAL, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    text = "独立完成“商析”全栈电商数据分析平台：基于 React + TypeScript 构建多看板可视化与 CSV 字段映射流程，使用 Django REST + MySQL 实现用户认证、数据集隔离、导入任务与关系确认；支持经营总览、渠道/商品分析、用户 RFM 分层、行为漏斗、复购趋势等分析能力。"
    r2 = p2.add_run(text)
    set_font(r2, size=10.5, color=INK)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.34)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    for name in ("List Bullet",):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("SHANGXI · PROJECT PORTFOLIO")
    set_font(r, size=8, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("商析 · 全栈电商数据分析平台")
    set_font(r, size=8, color=MUTED)


def main():
    doc = Document()
    configure(doc)

    # Cover - customer_story header pattern, adapted to a project case study.
    add_text(doc, "PROJECT PORTFOLIO", size=10, color=SIGNAL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=46, after=10)
    add_text(doc, "商析", size=30, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_text(doc, "全栈电商数据分析平台", size=16, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_text(doc, "从 CSV 多表导入、数据结构确认到经营分析看板的一体化作品", size=10.5, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    rule = doc.add_paragraph()
    add_rule(rule, color=SIGNAL, sz="14")
    add_text(doc, "项目定位", size=10.5, color=OLIVE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=4)
    add_text(doc, "面向电商运营场景的数据导入、数据治理与经营洞察原型", size=13, color=INK, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=26)
    add_metric_strip(doc)
    add_text(doc, "作品集说明：页面截图来自本地运行的项目实例；示例数据为模拟电商经营数据，不包含真实业务或个人隐私数据。", size=8.5, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=28, after=0)

    doc.add_page_break()
    add_heading(doc, "一、项目概述", 1)
    add_text(doc, "“商析”解决的是电商数据从零散 CSV 文件到可行动经营结论之间的断层：运营人员可一次导入订单、订单明细、商品、用户及行为表；系统识别字段并提示表关系风险，确认后进入经营、用户与行为分析看板。", size=10.5, after=10)
    add_heading(doc, "核心能力", 2)
    add_bullet(doc, "多表数据接入：", "支持订单、商品、用户、行为等 CSV 同时导入，在浏览器端完成解析、预览和初步校验。")
    add_bullet(doc, "字段治理：", "对“每行代表、单表类型、文件角色、标准字段映射”进行可视化确认，避免错误口径进入看板。")
    add_bullet(doc, "安全关系确认：", "识别一对一、一对多、多对一及高风险多对多关系，展示匹配率、唯一率和潜在数据膨胀风险。")
    add_bullet(doc, "经营分析：", "提供 GMV、订单、客单价、转化率、渠道贡献、商品机会等经营指标与建议。")
    add_bullet(doc, "用户与行为分析：", "提供用户属性、购买用户数、复购率、RFM 八类分层、行为漏斗、复购次数漏斗与趋势分析。")
    add_heading(doc, "技术架构", 2)
    add_architecture(doc)
    add_figure(doc, "overview.png", "图 1 经营总览：关键经营指标、销售趋势与当日经营简报。", width=6.35)

    doc.add_page_break()
    add_heading(doc, "二、数据导入与数据治理", 1)
    add_text(doc, "数据导入不是简单上传文件：系统将文件解析、角色识别、字段映射、数据质量和关系确认拆为明确步骤，使后续图表能追溯到字段与计算口径。", size=10.5, after=8)
    add_figure(doc, "import-queue-clean.png", "图 2 多文件导入队列：订单、明细、商品、用户、行为表均已完成浏览器端结构解析。", width=6.10)
    add_heading(doc, "字段与关系确认机制", 2)
    add_bullet(doc, "表级配置：", "为每个文件选择行粒度、单表类型与文件角色，支持订单主表、订单明细、商品表、用户表与用户行为表。")
    add_bullet(doc, "字段级映射：", "将源字段映射为用户 ID、订单 ID、商品 ID、时间、金额、行为、渠道、地址、性别、设备等标准字段。")
    add_bullet(doc, "数据质量：", "汇总总行数、空值、日期与金额问题，并说明不可用指标与补齐方式。")
    add_bullet(doc, "关系防护：", "只有确认后的关系进入正式分析；低唯一率、多对多关系会被标记为风险，避免错误关联导致 GMV 或用户数被重复放大。")
    add_figure(doc, "field-mapping.png", "图 3 数据关系核验：展示匹配率、右键唯一率及多对多风险提示。", width=6.15)

    doc.add_page_break()
    add_heading(doc, "三、看板与分析能力", 1)
    add_text(doc, "项目将数据能力落实到可操作的看板。每张图表独立判断数据字段是否可用：数据不足时仅隐藏对应图表，而不影响同页面其他分析模块。", size=10.5, after=10)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [1900, 3160, 4300])
    headers = ["模块", "核心指标 / 图表", "业务使用场景"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "E8E1D5")
        set_cell_margins(cell)
        set_cell_border(cell, top={"val":"single","sz":"6","color":LINE}, bottom={"val":"single","sz":"6","color":LINE}, left={"val":"single","sz":"6","color":LINE}, right={"val":"single","sz":"6","color":LINE})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        set_font(r, size=9.5, color=INK, bold=True)
    rows = [
        ("经营总览", "GMV、支付订单、客单价、转化率、渠道贡献、销售趋势", "快速判断增长、渠道差异与优先处理事项"),
        ("商品和品类", "商品排行、品类排行、销售额与销量", "识别高增长商品和低转化商品"),
        ("用户分析", "属性分布、购买用户数、复购率、RFM 八类分层", "做用户分层、会员维护与召回策略"),
        ("行为分析", "行为漏斗、复购次数漏斗、复购率 / 销售额趋势", "定位转化损耗，区分短期销售与复购质量"),
        ("数据助手", "基于当前数据上下文的分析问答入口", "帮助用户理解字段、指标和看板结论"),
    ]
    for index, values in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_shading(cell, PAPER if index % 2 == 0 else "F8F5EF")
            set_cell_margins(cell)
            set_cell_border(cell, top={"val":"single","sz":"4","color":LINE}, bottom={"val":"single","sz":"4","color":LINE}, left={"val":"single","sz":"4","color":LINE}, right={"val":"single","sz":"4","color":LINE})
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=9.2, color=INK, bold=(cell == cells[0]))
    add_text(doc, "注：上述能力会受实际上传字段、已确认关系和数据质量约束；系统在确认页说明缺失原因，而不是生成误导性图表。", size=8.5, color=MUTED, italic=True, before=5, after=10)
    add_figure(doc, "intake.png", "图 4 数据导入空状态：提示用户从订单、商品、用户文件开始构建可分析的数据集。", width=6.15)

    doc.add_page_break()
    add_heading(doc, "四、全栈实现与工程化要点", 1)
    add_heading(doc, "前端", 2)
    add_bullet(doc, "React + TypeScript：", "构建模块化看板、数据导入、字段确认、用户中心和数据助手页面。")
    add_bullet(doc, "浏览器端 CSV 处理：", "完成文件解析、字段推断、数据预览和关系诊断，降低无效上传与服务端压力。")
    add_bullet(doc, "图表与交互：", "支持趋势、饼图、柱状图、漏斗和 RFM 可视化；筛选条件统一影响关联指标。")
    add_heading(doc, "后端与数据层", 2)
    add_bullet(doc, "Django REST：", "实现注册、登录、退出、CSRF Token、Session Cookie 与数据权限接口。")
    add_bullet(doc, "MySQL：", "存储用户、数据集元信息、字段映射、关系确认、导入任务及标准化业务数据。")
    add_bullet(doc, "数据隔离：", "普通用户仅访问本人数据集和公共演示数据；公共数据集默认只读。")
    add_bullet(doc, "稳定性处理：", "处理 CORS / CSRF 跨端请求、重复数据集检测、导入批次限制与失败状态回收。")
    add_figure(doc, "user-center.png", "图 5 用户中心：展示账号状态、角色与数据访问范围。", width=6.10)
    add_resume_block(doc)
    add_text(doc, "项目源码结构：前端 src/，后端 backend/，示例数据 demo-data/；本地开发环境通过 Vite、Django 与 MySQL 协同运行。", size=8.8, color=MUTED, italic=True, before=14, after=0)

    doc.core_properties.title = "商析 - 全栈电商数据分析平台作品集"
    doc.core_properties.subject = "个人简历项目作品集"
    doc.core_properties.author = ""
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
