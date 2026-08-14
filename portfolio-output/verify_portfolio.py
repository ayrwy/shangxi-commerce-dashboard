from pathlib import Path
from docx import Document

path = Path(__file__).resolve().parent / "商析_全栈电商数据分析平台_作品集.docx"
document = Document(path)
print(f"paragraphs={len(document.paragraphs)}")
print(f"tables={len(document.tables)}")
print(f"inline_shapes={len(document.inline_shapes)}")
print(f"sections={len(document.sections)}")
print(f"title={document.core_properties.title}")
assert len(document.inline_shapes) == 5
assert len(document.tables) == 4
assert document.core_properties.title == "商析 - 全栈电商数据分析平台作品集"
